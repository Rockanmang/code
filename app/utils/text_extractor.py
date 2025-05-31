"""
文本提取工具函数
支持从PDF、DOCX、HTML等文件中提取文本和标题
使用PyMuPDF处理复杂格式的PDF文档
"""

import os
import re
from pathlib import Path
from typing import Optional, List, Dict
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_pdf_text_with_pymupdf(file_path: str) -> Optional[str]:
    """
    使用PyMuPDF从PDF文件中提取文本，支持复杂格式
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回None
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        
        # 提取所有页面的文本
        text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # 尝试多种文本提取方法
            # 方法1: 尝试按阅读顺序提取文本
            try:
                page_text = page.get_text(sort=True)
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    continue
            except:
                pass
            
            # 方法2: 使用字典格式提取，更精确控制
            try:
                blocks = page.get_text("dict")
                page_text = ""
                
                # 按块处理文本
                if "blocks" in blocks:
                    for block in blocks["blocks"]:
                        if "lines" in block:
                            for line in block["lines"]:
                                if "spans" in line:
                                    line_text = ""
                                    for span in line["spans"]:
                                        if "text" in span:
                                            line_text += span["text"]
                                    if line_text.strip():
                                        page_text += line_text + "\n"
                
                if page_text.strip():
                    text += page_text + "\n"
                    continue
            except:
                pass
            
            # 方法3: 简单文本提取作为后备
            try:
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
            except:
                logger.warning(f"页面 {page_num + 1} 文本提取失败")
        
        doc.close()
        
        # 清理文本
        text = clean_extracted_text(text)
        
        if text.strip():
            logger.info(f"使用PyMuPDF成功从PDF提取文本，长度: {len(text)} 字符")
            return text
        else:
            logger.warning(f"PyMuPDF从PDF文件 {file_path} 中没有提取到文本")
            return ""
            
    except ImportError:
        logger.warning("PyMuPDF库未安装，回退到PyPDF2")
        return None
    except Exception as e:
        logger.error(f"PyMuPDF提取PDF文本失败: {e}")
        return None

def extract_pdf_text(file_path: str) -> Optional[str]:
    """
    使用多种库从PDF文件中提取文本，优先使用PyMuPDF
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回None
    """
    # 首先尝试使用PyMuPDF
    text = extract_pdf_text_with_pymupdf(file_path)
    if text is not None:
        return text
    
    # 回退到PyPDF2
    logger.info("回退到PyPDF2进行PDF文本提取")
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # 提取所有页面的文本
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # 清理文本
            text = clean_extracted_text(text)
            
            if text.strip():
                logger.info(f"使用PyPDF2从PDF提取文本，长度: {len(text)} 字符")
                return text
            else:
                logger.warning(f"PDF文件 {file_path} 中没有可提取的文本")
                return ""
                
    except ImportError:
        logger.error("PyPDF2库未安装，无法提取PDF文本")
        return ""
    except Exception as e:
        logger.error(f"PyPDF2提取PDF文本失败: {e}")
        return ""

def extract_pdf_text_enhanced(file_path: str) -> Dict[str, any]:
    """
    增强的PDF文本提取，返回详细信息
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        Dict: 包含文本、页数、提取方法等信息
    """
    result = {
        "text": "",
        "page_count": 0,
        "extraction_method": None,
        "has_images": False,
        "has_tables": False,
        "text_blocks": [],
        "extraction_success": False
    }
    
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        result["page_count"] = len(doc)
        result["extraction_method"] = "PyMuPDF"
        
        all_text = ""
        text_blocks = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # 检查页面内容类型
            page_dict = page.get_text("dict")
            
            # 检测图像和表格
            if page.get_images():
                result["has_images"] = True
            
            # 提取文本块
            page_text = ""
            if "blocks" in page_dict:
                for block in page_dict["blocks"]:
                    if "lines" in block:  # 文本块
                        block_text = ""
                        for line in block["lines"]:
                            if "spans" in line:
                                line_text = ""
                                for span in line["spans"]:
                                    if "text" in span:
                                        line_text += span["text"]
                                if line_text.strip():
                                    block_text += line_text + "\n"
                        
                        if block_text.strip():
                            text_blocks.append({
                                "page": page_num + 1,
                                "text": block_text.strip(),
                                "bbox": block.get("bbox", None)
                            })
                            page_text += block_text
            
            all_text += page_text
        
        doc.close()
        
        # 清理文本
        all_text = clean_extracted_text(all_text)
        
        result["text"] = all_text
        result["text_blocks"] = text_blocks
        result["extraction_success"] = bool(all_text.strip())
        
        logger.info(f"增强PDF提取完成: {len(all_text)} 字符, {len(text_blocks)} 个文本块")
        
    except ImportError:
        logger.warning("PyMuPDF不可用，使用标准提取")
        text = extract_pdf_text(file_path)
        result["text"] = text or ""
        result["extraction_method"] = "PyPDF2"
        result["extraction_success"] = bool(text and text.strip())
    except Exception as e:
        logger.error(f"增强PDF提取失败: {e}")
        text = extract_pdf_text(file_path)
        result["text"] = text or ""
        result["extraction_method"] = "fallback"
        result["extraction_success"] = bool(text and text.strip())
    
    return result

def extract_docx_text(file_path: str) -> Optional[str]:
    """
    使用python-docx从DOCX文件中提取文本
    
    Args:
        file_path: DOCX文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回空字符串
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取所有段落的文本
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        # 清理文本
        text = clean_extracted_text(text)
        
        if text.strip():
            logger.info(f"成功从DOCX提取文本，长度: {len(text)} 字符")
            return text
        else:
            logger.warning(f"DOCX文件 {file_path} 中没有可提取的文本")
            return ""
            
    except ImportError:
        logger.error("python-docx库未安装，无法提取DOCX文本")
        return ""
    except Exception as e:
        logger.error(f"提取DOCX文本失败: {e}")
        return ""

def extract_html_text(file_path: str) -> Optional[str]:
    """
    使用BeautifulSoup从HTML文件中提取文本
    
    Args:
        file_path: HTML文件路径
        
    Returns:
        Optional[str]: 提取的文本内容，失败时返回空字符串
    """
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        # 解析HTML
        soup = BeautifulSoup(content, 'html.parser')
        
        # 移除script和style标签
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 提取文本
        text = soup.get_text()
        
        # 清理文本
        text = clean_extracted_text(text)
        
        if text.strip():
            logger.info(f"成功从HTML提取文本，长度: {len(text)} 字符")
            return text
        else:
            logger.warning(f"HTML文件 {file_path} 中没有可提取的文本")
            return ""
            
    except ImportError:
        logger.error("beautifulsoup4库未安装，无法提取HTML文本")
        return ""
    except Exception as e:
        logger.error(f"提取HTML文本失败: {e}")
        return ""

def clean_extracted_text(text: str) -> str:
    """
    清理提取的文本（去除多余空白、特殊字符等）
    
    Args:
        text: 原始文本
        
    Returns:
        str: 清理后的文本
    """
    if not text:
        return ""
    
    # 替换多个空白字符为单个空格
    text = re.sub(r'\s+', ' ', text)
    
    # 移除行首行尾空白
    text = text.strip()
    
    # 移除特殊控制字符
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    return text

def extract_title_from_text(text: str, max_length: int = 100) -> str:
    """
    从文本中提取标题（改进版）
    
    Args:
        text: 文本内容
        max_length: 标题最大长度
        
    Returns:
        str: 提取的标题
    """
    if not text or not text.strip():
        return "未知标题"
    
    # 清理文本
    text = text.strip()
    
    # 按行分割，寻找合适的标题行
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    if lines:
        # 寻找第一个看起来像标题的行（长度适中，不是纯数字）
        for line in lines[:5]:  # 只检查前5行
            if 10 <= len(line) and not line.isdigit():
                # 移除常见的标题前缀
                line = re.sub(r'^(第\d+章|Chapter\s+\d+|Abstract|摘要|引言|Introduction)[:：\s]*', '', line, flags=re.IGNORECASE)
                if line:
                    # 截断标题
                    if len(line) > max_length:
                        return line[:max_length]
                    return line
        
        # 如果没找到合适的，使用第一行（截断）
        if len(lines[0]) > max_length:
            return lines[0][:max_length]
        return lines[0]
    
    return "未知标题"

def extract_title_from_filename(filename: str) -> str:
    """
    从文件名中提取标题
    
    Args:
        filename: 文件名
        
    Returns:
        str: 提取的标题
    """
    # 获取文件名（不含扩展名）
    title = Path(filename).stem
    
    # 将下划线和连字符替换为空格
    title = title.replace('_', ' ').replace('-', ' ')
    
    # 清理多余空格
    title = ' '.join(title.split())
    
    return title

def extract_text_from_file(file_path: str) -> str:
    """
    根据文件类型提取文本内容
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容，失败时返回空字符串
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_pdf_text(file_path) or ""
    elif file_ext in ['.docx', '.doc']:
        return extract_docx_text(file_path) or ""
    elif file_ext in ['.html', '.htm']:
        return extract_html_text(file_path) or ""
    elif file_ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            return clean_extracted_text(text)
        except Exception as e:
            logger.error(f"读取TXT文件失败: {e}")
            return ""
    else:
        logger.warning(f"不支持的文件类型: {file_ext}")
        return ""

def extract_metadata_from_file(file_path: str, original_filename: str) -> dict:
    """
    从文件中提取元数据（标题等）
    
    Args:
        file_path: 文件路径
        original_filename: 原始文件名
        
    Returns:
        dict: 包含提取元数据的字典
    """
    file_ext = Path(file_path).suffix.lower()
    
    metadata = {
        "title": extract_title_from_filename(original_filename),
        "extracted_text": "",
        "extraction_success": False,
        "file_type": file_ext,
        "text_length": 0
    }
    
    try:
        # 提取文本内容
        extracted_text = extract_text_from_file(file_path)
        
        if extracted_text and extracted_text.strip():
            # 从提取的文本中获取更好的标题
            title_from_text = extract_title_from_text(extracted_text)
            if title_from_text and title_from_text != "未知标题":
                metadata["title"] = title_from_text
            
            metadata["extracted_text"] = extracted_text
            metadata["extraction_success"] = True
            metadata["text_length"] = len(extracted_text)
            
            logger.info(f"成功提取文件元数据: {metadata['title']} (长度: {metadata['text_length']})")
        else:
            logger.warning(f"文件 {file_path} 文本提取失败或为空，使用文件名作为标题")
            
    except Exception as e:
        logger.error(f"文件元数据提取失败: {e}")
    
    return metadata

def is_text_extractable(file_path: str) -> bool:
    """
    检查文件是否支持文本提取
    
    Args:
        file_path: 文件路径
        
    Returns:
        bool: 是否支持文本提取
    """
    file_ext = Path(file_path).suffix.lower()
    supported_types = ['.pdf', '.docx', '.doc', '.html', '.htm', '.txt']
    return file_ext in supported_types

def estimate_reading_time(text: str, words_per_minute: int = 200) -> int:
    """
    估算文本阅读时间（分钟）
    
    Args:
        text: 文本内容
        words_per_minute: 每分钟阅读词数
        
    Returns:
        int: 估算的阅读时间（分钟）
    """
    if not text:
        return 0
    
    # 简单的词数统计（中英文混合）
    words = len(text.split())
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    
    # 中文字符按2个字符算一个词
    total_words = words + chinese_chars // 2
    
    reading_time = max(1, total_words // words_per_minute)
    return reading_time